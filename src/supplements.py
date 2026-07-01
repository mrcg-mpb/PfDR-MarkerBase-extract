"""
Supplement loader — turn the diverse files in a paper's supplement folder into
content blocks the model can actually read.

The MAIN paper is always a PDF and is handled directly by the agents. This module
is ONLY for SUPPLEMENTARY files, which arrive in many scientific formats. Each
file is converted to an Anthropic content block:

  PDF                         -> document block (passed through natively)
  XLSX / XLS (spreadsheets)   -> text, one tab-separated table per sheet
  CSV / TSV / TXT / tab-delim -> text (passed through as-is)
  DOCX / DOC (Word)           -> extracted text
  anything else               -> skipped, recorded in the summary

Deterministic, no LLM. Invoked ONLY on the supplement path (an AWAIT_SUPPLEMENT
re-check in stage 1, or an eligible paper's extraction in stage 2), so the common
no-supplement case keeps its exact previous code and cost.

Conversions are best-effort: a file we can't read is skipped with a note rather
than aborting the run. Spreadsheet/Word support needs the libraries in
requirements.txt (openpyxl, xlrd, python-docx); legacy .doc additionally needs
LibreOffice (`soffice`, present on the CI runner) or antiword/catdoc.
"""
import base64
import io
import os
import shutil
import subprocess
import tempfile

MAX_CHARS_PER_FILE = 60_000     # keep token cost bounded; note when we truncate
MAX_ROWS_PER_SHEET = 1000       # per spreadsheet sheet

# Extensions we render as plain text by passing the raw contents straight through.
PLAINTEXT_EXTS = {"csv", "tsv", "tab", "txt", "text", ""}


def load(files):
    """files: [{'name': str, 'bytes': bytes}, ...]. Returns (blocks, summary).

    `blocks` is a list of Anthropic content blocks (document/text) ready to drop
    into a message; `summary` is a short human-readable line of what was loaded
    or skipped (for logging).
    """
    blocks, notes = [], []
    for f in files or []:
        name = f.get("name", "file")
        data = f.get("bytes", b"")
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        try:
            if ext == "pdf":
                blocks.append(_pdf_block(data))
                notes.append(f"{name} (pdf)")
                continue
            if ext == "xlsx":
                text, extra = _xlsx_to_text(data)
            elif ext == "xls":
                text, extra = _xls_to_text(data)
            elif ext == "docx":
                text, extra = _docx_to_text(data)
            elif ext == "doc":
                text, extra = _doc_to_text(data)
            elif ext in PLAINTEXT_EXTS:
                text, extra = _read_text(data), ""
            else:
                notes.append(f"{name} (skipped: unsupported .{ext})")
                continue
            text, trunc = _truncate(text)
            blocks.append(_text_block(name, text))
            notes.append(f"{name} ({ext or 'text'}{extra}{trunc})")
        except Exception as e:  # never let one bad file abort the run
            notes.append(f"{name} (unreadable: {e})")
    return blocks, ("; ".join(notes) if notes else "no files")


# --- Content blocks --------------------------------------------------------

def _pdf_block(data):
    b64 = base64.standard_b64encode(data).decode("utf-8")
    return {"type": "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": b64}}


def _text_block(name, text):
    return {"type": "text", "text": f"----- Supplementary file: {name} -----\n{text}"}


# --- Format converters -----------------------------------------------------

def _read_text(data):
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _truncate(text):
    if len(text) <= MAX_CHARS_PER_FILE:
        return text, ""
    return text[:MAX_CHARS_PER_FILE] + "\n... [truncated]", ", truncated"


def _xlsx_to_text(data):
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts, truncated = [], False
    try:
        for ws in wb.worksheets:
            parts.append(f"# Sheet: {ws.title}")
            n = 0
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if not any(c.strip() for c in cells):
                    continue
                parts.append("\t".join(cells))
                n += 1
                if n >= MAX_ROWS_PER_SHEET:
                    parts.append(f"... [sheet truncated at {MAX_ROWS_PER_SHEET} rows]")
                    truncated = True
                    break
    finally:
        wb.close()
    return "\n".join(parts), (", rows truncated" if truncated else "")


def _xls_to_text(data):
    import xlrd
    wb = xlrd.open_workbook(file_contents=data)
    parts, truncated = [], False
    for sh in wb.sheets():
        parts.append(f"# Sheet: {sh.name}")
        for r in range(min(sh.nrows, MAX_ROWS_PER_SHEET)):
            parts.append("\t".join(str(sh.cell_value(r, c)) for c in range(sh.ncols)))
        if sh.nrows > MAX_ROWS_PER_SHEET:
            parts.append(f"... [sheet truncated at {MAX_ROWS_PER_SHEET} rows]")
            truncated = True
    return "\n".join(parts), (", rows truncated" if truncated else "")


def _docx_to_text(data):
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts), ""


def _doc_to_text(data):
    """Legacy binary .doc: convert via LibreOffice (present on the CI runner and
    macOS), falling back to antiword/catdoc. Raises if no converter is available
    (load() turns that into a skip-with-note)."""
    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "in.doc")
        with open(src, "wb") as fh:
            fh.write(data)
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            # Isolated per-call profile dir so concurrent/repeated headless runs
            # don't clash on the default LibreOffice profile lock.
            subprocess.run([soffice, "--headless",
                            f"-env:UserInstallation=file://{os.path.join(tmp, 'loprofile')}",
                            "--convert-to", "txt:Text", "--outdir", tmp, src],
                           capture_output=True, timeout=120)
            out = os.path.join(tmp, "in.txt")
            if os.path.exists(out):
                with open(out, "rb") as fh:
                    return _read_text(fh.read()), ""
        for tool in ("antiword", "catdoc"):
            if shutil.which(tool):
                p = subprocess.run([tool, src], capture_output=True, timeout=60)
                if p.returncode == 0 and p.stdout:
                    return _read_text(p.stdout), ""
        raise RuntimeError("no .doc converter available (need libreoffice/antiword)")
